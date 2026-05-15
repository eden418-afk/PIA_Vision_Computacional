from __future__ import annotations

import re
import zipfile
from pathlib import Path

import cv2
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parents[2]
TMP_DIR = BASE_DIR / "tmp" / "docs"
OUT_DIR = BASE_DIR / "output" / "doc"
REPORT_MD = BASE_DIR / "reporte_pia_vision_computacional.md"

THUMB_DIR = TMP_DIR / "freecad_thumbnails"
PROCESSING_DIR = TMP_DIR / "generated_images"


def ensure_dirs() -> None:
    THUMB_DIR.mkdir(parents=True, exist_ok=True)
    PROCESSING_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)


def extract_thumbnail(fcstd_name: str, out_name: str) -> Path:
    fcstd_path = BASE_DIR / "cad" / "modelos" / fcstd_name
    out_path = THUMB_DIR / out_name
    with zipfile.ZipFile(fcstd_path, "r") as zf:
        data = zf.read("thumbnails/Thumbnail.png")
    out_path.write_bytes(data)
    return out_path


def build_processing_montage() -> Path:
    original = cv2.imread(str(BASE_DIR / "dataset" / "originales" / "pieza_05.png"))
    result = cv2.imread(str(BASE_DIR / "resultados" / "imagenes_resultado" / "resultado_pieza_05.png"))
    if original is None or result is None:
        raise FileNotFoundError("No se pudieron cargar las imágenes para el montaje.")

    target_h = 500
    scale_o = target_h / original.shape[0]
    scale_r = target_h / result.shape[0]
    original = cv2.resize(original, (int(original.shape[1] * scale_o), target_h))
    result = cv2.resize(result, (int(result.shape[1] * scale_r), target_h))

    gap = 40
    width = original.shape[1] + result.shape[1] + gap * 3
    height = target_h + 120
    canvas = 255 * cv2.UMat(height, width, cv2.CV_8UC3).get()

    x1 = gap
    x2 = original.shape[1] + gap * 2
    y = 80
    canvas[y:y + target_h, x1:x1 + original.shape[1]] = original
    canvas[y:y + target_h, x2:x2 + result.shape[1]] = result

    cv2.putText(canvas, "Imagen sintetica original", (x1, 45), cv2.FONT_HERSHEY_SIMPLEX, 1.0, (0, 0, 0), 2)
    cv2.putText(canvas, "Resultado procesado con OpenCV", (x2, 45), cv2.FONT_HERSHEY_SIMPLEX, 1.0, (0, 0, 0), 2)

    out_path = PROCESSING_DIR / "secuencia_procesamiento.png"
    cv2.imwrite(str(out_path), canvas)
    return out_path


def preprocess_markdown(text: str, image_map: dict[str, Path]) -> str:
    replacements = {
        "![Ejemplo de pieza CAD con centroide y eje angular](ruta/a/captura_freecad.png)":
        f"![Ejemplo de pieza CAD con centroide y eje angular]({image_map['freecad_main'].as_posix()})",
        "![Secuencia de procesamiento OpenCV](ruta/a/diagrama_procesamiento.png)":
        f"![Secuencia de procesamiento OpenCV]({image_map['processing'].as_posix()})",
        "![Captura adicional de FreeCAD 1](ruta/a/captura_freecad_01.png)":
        f"![Captura adicional de FreeCAD 1]({image_map['freecad_extra_1'].as_posix()})",
        "![Captura adicional de FreeCAD 2](ruta/a/captura_freecad_02.png)":
        f"![Captura adicional de FreeCAD 2]({image_map['freecad_extra_2'].as_posix()})",
        "![Hoja de MathCAD o captura equivalente](ruta/a/hoja_mathcad.png)":
        "[pendiente] Captura de hoja de MathCAD o cálculo externo no incluida en el repositorio.",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_default_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_runs_from_markdown(paragraph, text: str) -> None:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_image(doc: Document, path_str: str, caption: str) -> None:
    path = Path(path_str)
    if not path.is_absolute():
        path = (BASE_DIR / path).resolve()

    if not path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[pendiente] {caption}")
        run.italic = True
        return

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(path), width=Inches(6.2))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_cap.add_run(caption)
    run.italic = True


def add_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if set(line.strip()) <= {"|", "-", ":", " "}:
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def build_doc(markdown_text: str) -> Document:
    doc = Document()
    set_default_styles(doc)

    lines = markdown_text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                for idx, code_line in enumerate(code_lines):
                    run = p.add_run(code_line + ("\n" if idx < len(code_lines) - 1 else ""))
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        img_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if img_match:
            add_image(doc, img_match.group(2), img_match.group(1))
            i += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs_from_markdown(p, text)
            i += 1
            continue

        list_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if list_match:
            p = doc.add_paragraph(style="List Number")
            add_runs_from_markdown(p, list_match.group(1))
            i += 1
            continue

        bullet_match = re.match(r"^-+\s+(.*)$", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_from_markdown(p, bullet_match.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        if stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs_from_markdown(p, stripped)
        i += 1

    return doc


def main() -> None:
    ensure_dirs()

    image_map = {
        "freecad_main": extract_thumbnail("pieza_01.FCStd", "pieza_01_thumb.png"),
        "freecad_extra_1": extract_thumbnail("pieza_05.FCStd", "pieza_05_thumb.png"),
        "freecad_extra_2": extract_thumbnail("pieza_10.FCStd", "pieza_10_thumb.png"),
        "processing": build_processing_montage(),
    }

    md_text = REPORT_MD.read_text(encoding="utf-8")
    md_text = preprocess_markdown(md_text, image_map)

    doc = build_doc(md_text)

    out_docx = OUT_DIR / "reporte_pia_vision_computacional.docx"
    doc.save(out_docx)

    root_copy = BASE_DIR / "reporte_pia_vision_computacional.docx"
    doc.save(root_copy)

    print(f"DOCX generado: {out_docx}")
    print(f"Copia generada: {root_copy}")


if __name__ == "__main__":
    main()
